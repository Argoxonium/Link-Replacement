from docx import Document
import csv
import re
import requests

def main():
    # Example usage
    hyperlinks = extract_hyperlinks(r"C:\Users\nhorn\Documents\SRC\Notes, Documenting, and Ect\Work Planner Aid_Rev8.docx")
    for link in hyperlinks:
        print(f"Text: {link['text']}, URL: {link['url']}")

    # Example usage
    doc_path = r"C:\Users\nhorn\Documents\SRC\Notes, Documenting, and Ect\Work Planner Aid_Rev8.docx"
    csv_path = r"C:\Users\nhorn\Documents\SRC\Notes, Documenting, and Ect\Work Planner Aid_Rev8 links.csv"
    hyperlinks = extract_hyperlinks(doc_path)
    save_to_csv(hyperlinks, csv_path)

    print(f"Hyperlinks have been written to {csv_path}.")

"""
This functions goal is to identify the links attached to specific text within a word document. This test will be used later to help
Identify what links are needed useing the new updated URL.
"""
def extract_hyperlinks(doc_path:str):
    # Load the document
    doc = Document(doc_path)
    hyperlinks:list = []

    # Access the document's relationship part for hyperlinks
    rels = doc.part.rels

    # Loop through each relationship
    for rel in rels.values():
        if "hyperlink" in rel.reltype:  # Check if it is a hyperlink
            link_info:dict = {'text': None, 'url': rel.target_ref}
            
            # Search for the text associated with this hyperlink
            for paragraph in doc.paragraphs:
                if rel.rId in paragraph._p.xml:
                    link_info['text'] = paragraph.text
                    break
            
            hyperlinks.append(link_info)

    return hyperlinks

'''
This takes the links that were found and saves them to a csv file.
'''

def save_to_csv(hyperlink:str, csv_path:str):
    # Write the hyperlinks to a CSV file
    with open(csv_path, mode='w', newline='', encoding='utf-8') as file:
        writer = csv.writer(file)
        writer.writerow(['Text', 'URL'])  # Write header
        for link in hyperlink:
            writer.writerow([link['text'], link['url']])

#Extracting the important information from the link to determine what needs to be updated
def pull_relative_info(url:str)->str:
    #use regular expression t find the last segment after the last '/'
    match = re.search(r'/([^/]+)\.pdf$', url)

    #extract the information without the pdf. If no information return none.
    if match:
        return match.group(1)
    else:
        return None

#updateing the link string. 
def create_new_link(updated_link:str, links:dict) -> dict:
    #create a new dict
    links_transformed:dict = {}

    #loop through each item in my dict and pull information needed and create a new dict compairing new and old links.
    for text, link in links.items():
        link_piece = pull_relative_info(link) #pull info needed
        new_link = f"{updated_link}{link_piece}" #create new link
        links_transformed[link] = new_link #add to dict

    return links_transformed

def update_link(doc_path: str, links: dict) -> None:
    # Load the document
    doc = Document(doc_path)
    
    # Access the document's relationship part for hyperlinks
    rels = doc.part.rels

    # Loop through each paragraph and check for hyperlinks
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.hyperlink:  # Check if the run contains a hyperlink
                # Get the hyperlink relationship ID
                hyperlink_id = run.hyperlink.rId
                
                # If the hyperlink exists in the dictionary, update it
                if hyperlink_id in rels and rels[hyperlink_id].target_ref in links:
                    new_url = links[rels[hyperlink_id].target_ref]
                    rels[hyperlink_id]._target = new_url


    # Save the updated document
    doc.save('updated_document.docx')

#test if the link works
#TODO Update to accept redirects
def check_link(url):
    try:
        response = requests.head(url, allow_redirects=True)
        # Check if the status code is 200 (OK)
        if response.status_code == 200:
            return True
        else:
            return False
    except requests.RequestException as e:
        print(f"Error checking URL {url}: {e}")
        return False

if __name__ == "__main__":
    main()